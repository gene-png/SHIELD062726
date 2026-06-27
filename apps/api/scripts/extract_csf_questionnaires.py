"""Extract CSF tier interview questionnaires from Kentro .docx files.

Reads (from the repo's reference-docs/ folder):
  reference-docs/Step_1_1_Interview_HIGH.docx
  reference-docs/Step_1_2_Interview_MODERATE.docx
  reference-docs/Step_1_3_Interview_LOW.docx

Writes (one file per tier):
  packages/csf-data/source/csf_tier_high.json
  packages/csf-data/source/csf_tier_moderate.json
  packages/csf-data/source/csf_tier_low.json

Output JSON shape mirrors packages/zt-data so the same loader pattern works
for both frameworks.

Run it (python-docx is a dev dependency):
  cd apps/api && python -m scripts.extract_csf_questionnaires

Or as a one-shot via docker without touching the local env:
  docker run --rm -v ${PWD}:/workspace -w /workspace python:3.12-slim \
    sh -c "pip install -q python-docx && \
      python apps/api/scripts/extract_csf_questionnaires.py"
"""

from __future__ import annotations

import json
import re
from collections.abc import Iterable
from dataclasses import asdict, dataclass, field
from pathlib import Path

from docx import Document
from docx.table import Table, _Cell

WORKSPACE = Path(__file__).resolve().parents[3]
SRC_DIR = WORKSPACE / "reference-docs"
OUT_DIR = WORKSPACE / "packages" / "csf-data" / "source"
OUT_DIR.mkdir(parents=True, exist_ok=True)

TIER_FILES: dict[str, tuple[str, str]] = {
    "high": ("Step_1_1_Interview_HIGH.docx", "csf-tier-high"),
    "moderate": ("Step_1_2_Interview_MODERATE.docx", "csf-tier-moderate"),
    "low": ("Step_1_3_Interview_LOW.docx", "csf-tier-low"),
}

SECTION_HEADER_RE = re.compile(r"^\s*(\d+)\.\s+(.+?)\s*$")
COLUMN_HEADER_TOKENS = {"#", "Question & Interviewer Cues", "CSF 2.0 Subcategories", "Response"}


@dataclass
class CsfQuestion:
    external_id: str
    section_id: str
    section_name: str
    order_in_section: int
    order_index: int
    stem: str
    cues: list[str] = field(default_factory=list)
    dimensions: str = ""
    csf_subcategories: list[str] = field(default_factory=list)
    implementation_groups: list[str] = field(default_factory=list)


def _row_text(row) -> list[str]:
    """Return per-cell text, but de-duplicate the merged-cell pattern python-docx exposes."""
    seen: list[str] = []
    for cell in row.cells:
        text = cell.text.strip()
        # python-docx repeats merged cells across positions; skip exact-adjacent dupes.
        if seen and seen[-1] == text:
            continue
        seen.append(text)
    return seen


def _is_section_header_row(row) -> tuple[str, str] | None:
    """Section header rows are all-merged into one cell like '1. Governance & Oversight'."""
    cells = _row_text(row)
    if len(cells) != 1:
        return None
    m = SECTION_HEADER_RE.match(cells[0])
    if not m:
        return None
    return m.group(1), m.group(2).strip()


def _is_column_header_row(row) -> bool:
    cells = set(_row_text(row))
    return bool(cells & COLUMN_HEADER_TOKENS) and "#" in cells


def _split_stem_and_cues(cell: _Cell) -> tuple[str, list[str]]:
    """First non-empty paragraph = stem; subsequent non-empty paragraphs = cues.

    The Kentro template puts the interviewer cues as italicized follow-ups under
    the main question. They occasionally span multiple paragraphs.
    """
    paras = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    if not paras:
        return "", []
    stem = paras[0]
    cues = paras[1:]
    return stem, cues


def _split_dimensions(cues: list[str]) -> tuple[list[str], str]:
    """The Kentro template trails each question with a 'Dimensions: A → B → C' tag.

    That's question metadata, not an interviewer cue — pull it out into its own
    field so the UI doesn't show it in the cue list.
    """
    if cues and cues[-1].lower().startswith("dimensions:"):
        return cues[:-1], cues[-1].split(":", 1)[1].strip()
    return cues, ""


def _parse_subcategories(text: str) -> tuple[list[str], list[str]]:
    """Cell contains CSF 2.0 subcategory IDs and optionally 'IG: 1, 2, 3' implementation groups."""
    text = text.replace("⏎", "\n").replace("\r", "\n")
    subcats: list[str] = []
    igs: list[str] = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.upper().startswith("IG:"):
            # Pull comma-separated digit groups out.
            igs.extend(part.strip() for part in line.split(":", 1)[1].split(",") if part.strip())
            continue
        # CSF Subcategory pattern: TWO_OR_THREE_LETTERS.TWO_LETTERS-NN
        for token in re.findall(r"[A-Z]{2,3}\.[A-Z]{2}-\d{2}", line):
            subcats.append(token)
    # Dedupe preserving order.
    seen: set[str] = set()
    deduped: list[str] = []
    for s in subcats:
        if s not in seen:
            seen.add(s)
            deduped.append(s)
    return deduped, igs


def _walk_question_tables(tables: Iterable[Table]) -> list[CsfQuestion]:
    """Walk every table in the doc; emit questions in document order.

    A single table may contain one or multiple sections (the MOD doc bundles
    several sections into one giant table). We track the "current section" as
    we walk rows.
    """
    out: list[CsfQuestion] = []
    overall_order = 0
    current_section_id: str | None = None
    current_section_name: str | None = None
    order_in_section = 0

    for table in tables:
        # Skip non-question tables: header row contains "Document" or "Field" instead of "#".
        header_cells = set(_row_text(table.rows[0])) if table.rows else set()
        if "Document" in header_cells or "Field" in header_cells:
            continue

        for row in table.rows:
            sec = _is_section_header_row(row)
            if sec is not None:
                current_section_id, current_section_name = sec
                order_in_section = 0
                continue
            if _is_column_header_row(row):
                continue
            cells_raw = list(row.cells)
            texts = _row_text(row)
            # A question row starts with a numeric "#" cell.
            if not texts or not texts[0].strip().isdigit():
                continue
            if current_section_id is None:
                # Defensive: question outside any section — skip.
                continue
            # Map back to actual cells (after dedupe) for paragraph access on the stem cell.
            # cells_raw still includes merged duplicates; find the cell whose text matches texts[1].
            stem_cell = None
            for cell in cells_raw:
                if cell.text.strip() == texts[1]:
                    stem_cell = cell
                    break
            stem, cues = (
                _split_stem_and_cues(stem_cell) if stem_cell is not None else (texts[1], [])
            )
            cues, dimensions = _split_dimensions(cues)
            subcat_text = texts[2] if len(texts) > 2 else ""
            subcats, igs = _parse_subcategories(subcat_text)

            order_in_section += 1
            overall_order += 1
            external_id = f"CSF-{current_section_id}-{order_in_section}"
            out.append(
                CsfQuestion(
                    external_id=external_id,
                    section_id=current_section_id,
                    section_name=current_section_name or "",
                    order_in_section=order_in_section,
                    order_index=overall_order,
                    stem=stem,
                    cues=cues,
                    dimensions=dimensions,
                    csf_subcategories=subcats,
                    implementation_groups=igs,
                )
            )
    return out


def extract_tier(tier: str, filename: str, framework_key: str) -> dict:
    path = SRC_DIR / filename
    doc = Document(str(path))
    questions = _walk_question_tables(doc.tables)
    sections: dict[str, dict] = {}
    for q in questions:
        sections.setdefault(
            q.section_id,
            {"section_id": q.section_id, "section_name": q.section_name, "question_count": 0},
        )
        sections[q.section_id]["question_count"] += 1
    return {
        "framework_key": framework_key,
        "tier": tier,
        "tier_label": {"high": "HIGH", "moderate": "MODERATE", "low": "LOW"}[tier],
        "source_file": filename,
        "maturity_levels": [
            "tier1-partial",
            "tier2-risk-informed",
            "tier3-repeatable",
            "tier4-adaptive",
        ],
        "maturity_descriptions": {
            "tier1-partial": "Cybersecurity is reactive and ad-hoc. Risk decisions aren't tied to a documented strategy and awareness is limited.",
            "tier2-risk-informed": "Risk decisions are informed by management but not consistently applied across the organization.",
            "tier3-repeatable": "Formal policies are in place and practiced consistently. The organization adapts to changes in threat and tech.",
            "tier4-adaptive": "Cybersecurity practices adapt continuously from lessons learned and predictive indicators; integrated into the org's culture.",
        },
        "sections": list(sections.values()),
        "questions": [asdict(q) for q in questions],
    }


def main() -> None:
    summary: list[str] = []
    for tier, (filename, framework_key) in TIER_FILES.items():
        data = extract_tier(tier, filename, framework_key)
        out_path = OUT_DIR / f"csf_tier_{tier}.json"
        out_path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
        summary.append(
            f"  {tier:9s}  {len(data['sections'])} sections  {len(data['questions'])} questions  -> {out_path.relative_to(WORKSPACE)}"
        )
    print("Extracted CSF tier questionnaires:")
    print("\n".join(summary))


if __name__ == "__main__":
    main()
