"""Shared python-docx helpers for service Word deliverables (Work Order C4).

Pure helpers over python-docx so each service's `render_docx` mirrors its PDF
(title, client, summary paragraphs, a findings table). python-docx is imported
lazily inside `new_document` so importing the exporter module stays cheap.
"""

from __future__ import annotations

import contextlib
import io
from collections.abc import Iterable, Sequence
from typing import Any


def new_document(title: str, *, author: str = "SHIELD by Kentro") -> Any:
    from docx import Document

    doc = Document()
    props = doc.core_properties
    props.title = title
    props.author = author
    return doc


def add_title(doc: Any, title: str, subtitle: str | None = None) -> None:
    doc.add_heading(title, level=0)
    if subtitle:
        doc.add_paragraph(subtitle)


def add_heading(doc: Any, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_page_break(doc: Any) -> None:
    doc.add_page_break()


def shade_cell(cell: Any, hex_color: str) -> None:
    """Set a table cell's background fill (python-docx has no direct API)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def add_paragraphs(doc: Any, lines: Iterable[str]) -> None:
    for line in lines:
        doc.add_paragraph(line)


def add_table(doc: Any, headers: Sequence[str], rows: Iterable[Sequence[Any]]) -> Any:
    """Add a styled table. Returns the table so callers can shade cells."""
    table = doc.add_table(rows=1, cols=len(headers))
    # Fall back to the default style if the template lacks the named one.
    with contextlib.suppress(KeyError):
        table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if val is None else str(val)
    return table


def to_bytes(doc: Any) -> bytes:
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# MIME type + extension for Word deliverables.
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
